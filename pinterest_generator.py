"""
Pinterest Pin Pack Generator
Uses Groq API to generate 50 Pinterest pin titles + descriptions
and saves them as a professional Word document.

Usage:
    python pinterest_generator.py
"""

import os
import json
from datetime import datetime
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─────────────────────────────────────────────
#  CONFIG — edit these before running
# ─────────────────────────────────────────────
GROQ_API_KEY = "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"   # get free at console.groq.com
NICHE        = "home organization"         # change to your client's niche
CLIENT_NAME  = "Sarah"                     # optional, used in doc header
NUM_PINS     = 50                          # keep at 50 for standard gig
# ─────────────────────────────────────────────


def generate_pins(niche: str, num_pins: int) -> list[dict]:
    """Call Groq API and get pin data as a list of dicts."""

    client = Groq(api_key=GROQ_API_KEY)

    prompt = f"""You are an expert Pinterest content strategist.
Generate exactly {num_pins} Pinterest pin ideas for the niche: "{niche}".

Rules:
- Each pin must have a TITLE (max 100 chars, punchy, keyword-rich)
- Each pin must have a DESCRIPTION (150-200 chars, includes a call to action, 3-5 hashtags)
- Vary the content types: tips, how-tos, lists, quotes, product ideas, before/after, inspirational
- Make titles SEO-friendly for Pinterest search
- Descriptions must sound natural, not robotic
- CRITICAL: Ensure MAXIMUM DIVERSITY across all pins. Do NOT repeat the same phrasing, ideas, or sentence structures.
- CRITICAL: Do NOT repeat the same hashtags endlessly. Use highly specific, varied hashtags for each individual pin.

Return ONLY a valid JSON array, no extra text. Format:
[
  {{
    "pin_number": 1,
    "title": "...",
    "description": "...",
    "content_type": "tip|how-to|list|quote|product|inspiration"
  }},
  ...
]"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.8,
        max_tokens=8000,
    )

    raw = response.choices[0].message.content.strip()

    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    pins = json.loads(raw)
    return pins


def save_to_docx(pins: list[dict], niche: str, client_name: str, output_path: str):
    """Save the pin list to a beautifully formatted Word document."""

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Brand colors (Pinterest red palette) ──
    PINTEREST_RED  = RGBColor(0xE6, 0x00, 0x23)
    DARK_GRAY      = RGBColor(0x33, 0x33, 0x33)
    MID_GRAY       = RGBColor(0x77, 0x77, 0x77)
    LIGHT_BG       = RGBColor(0xFF, 0xF5, 0xF5)

    # ── Header ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Pinterest Pin Pack")
    run.bold      = True
    run.font.size = Pt(26)
    run.font.color.rgb = PINTEREST_RED

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f'Niche: {niche.title()}')
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = MID_GRAY

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Prepared for: {client_name}   |   "
        f"Total pins: {len(pins)}   |   "
        f"Date: {datetime.now().strftime('%B %d, %Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = MID_GRAY

    doc.add_paragraph()  # spacer

    # ── Divider ──
    div = doc.add_paragraph("─" * 60)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.runs[0].font.color.rgb = PINTEREST_RED

    doc.add_paragraph()

    # ── Usage instructions box ──
    usage_para = doc.add_paragraph()
    usage_run = usage_para.add_run("HOW TO USE THIS PACK")
    usage_run.bold = True
    usage_run.font.size = Pt(11)
    usage_run.font.color.rgb = PINTEREST_RED

    instructions = [
        "Copy each title into your Pinterest pin title field.",
        "Copy the description into the pin description field.",
        "Add your image, then publish or schedule using Tailwind/Later.",
        "Post 3–5 pins/day for best reach. Mix content types.",
        "Add your website link to every pin!",
    ]
    for inst in instructions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(inst).font.size = Pt(10)

    doc.add_paragraph()

    # ── Pins ──
    content_type_colors = {
        "tip":         RGBColor(0x1D, 0x9E, 0x75),
        "how-to":      RGBColor(0x37, 0x8A, 0xDD),
        "list":        RGBColor(0xD8, 0x5A, 0x30),
        "quote":       RGBColor(0x7F, 0x77, 0xDD),
        "product":     RGBColor(0xBA, 0x75, 0x17),
        "inspiration": RGBColor(0xD4, 0x53, 0x7E),
    }

    for pin in pins:
        num     = pin.get("pin_number", "")
        title   = pin.get("title", "")
        desc    = pin.get("description", "")
        ctype   = pin.get("content_type", "tip").lower()
        color   = content_type_colors.get(ctype, PINTEREST_RED)

        # Pin number + type badge
        header_para = doc.add_paragraph()
        num_run = header_para.add_run(f"Pin {num:02d}  ")
        num_run.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = DARK_GRAY

        badge_run = header_para.add_run(f"[{ctype.upper()}]")
        badge_run.bold = True
        badge_run.font.size = Pt(9)
        badge_run.font.color.rgb = color

        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = DARK_GRAY
        title_p.paragraph_format.space_before = Pt(2)

        # Description
        desc_p = doc.add_paragraph()
        desc_run = desc_p.add_run(desc)
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = MID_GRAY
        desc_p.paragraph_format.space_after = Pt(8)

        # Thin separator every 5 pins
        if num % 5 == 0 and num < len(pins):
            sep = doc.add_paragraph("· " * 30)
            sep.runs[0].font.size = Pt(8)
            sep.runs[0].font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
            sep.paragraph_format.space_before = Pt(2)
            sep.paragraph_format.space_after  = Pt(2)

    # ── Footer note ──
    doc.add_paragraph()
    footer_div = doc.add_paragraph("─" * 60)
    footer_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_div.runs[0].font.color.rgb = MID_GRAY

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run(
        "This pack was custom-generated for your niche.\n"
        "For revisions or a fresh pack, contact your seller on Fiverr."
    )
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = MID_GRAY
    f_run.italic = True

    doc.save(output_path)
    print(f"  Saved: {output_path}")


def main():
    print("\n Pinterest Pin Pack Generator")
    print("=" * 40)

    niche       = input(f"Enter niche [{NICHE}]: ").strip() or NICHE
    client_name = input(f"Client name [{CLIENT_NAME}]: ").strip() or CLIENT_NAME

    print(f"\n  Calling Groq API for {NUM_PINS} pins in '{niche}'...")

    try:
        pins = generate_pins(niche, NUM_PINS)
        print(f"  Generated {len(pins)} pins successfully!")
    except json.JSONDecodeError as e:
        print(f"  JSON parse error: {e}")
        raise

    safe_niche  = niche.lower().replace(" ", "_")
    output_file = f"pinterest_pack_{safe_niche}_{datetime.now().strftime('%Y%m%d')}.docx"
    output_path = os.path.join(os.getcwd(), output_file)

    print(f"\n  Creating Word document...")
    save_to_docx(pins, niche, client_name, output_path)

    print(f"\n  Done! File ready: {output_file}")
    print(f"  Deliver this file to your Fiverr client.\n")


if __name__ == "__main__":
    main()
