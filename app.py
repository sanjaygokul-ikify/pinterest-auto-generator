"""
Pinterest Pin Pack Generator — Streamlit Web App
Clients enter their niche, click generate, download the Word doc.

Run locally:   streamlit run app.py
Deploy free:   streamlit.io (connect GitHub repo)
"""

import os
import json
import io
from datetime import datetime

import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─────────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Pinterest Pin Pack Generator",
    page_icon="📌",
    layout="centered",
)

# ─────────────────────────────────────────────
#  CUSTOM CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
    .main { background: #fff5f5; }
    h1 { color: #E60023 !important; }
    .stButton > button {
        background-color: #E60023;
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
        padding: 0.6rem 2rem;
        font-size: 16px;
    }
    .stButton > button:hover { background-color: #ad081b; color: white; }
    .metric-card {
        background: white;
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
        border: 1px solid #ffe0e0;
    }
    .download-box {
        background: #f0fff4;
        border: 1px solid #68d391;
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
    }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  GROQ API KEY — set via Streamlit secrets or env
# ─────────────────────────────────────────────
def get_api_key() -> str:
    # 1. Streamlit secrets (for deployed app) — add in Streamlit Cloud dashboard
    if "GROQ_API_KEY" in st.secrets:
        return st.secrets["GROQ_API_KEY"]
    # 2. Environment variable (for local run)
    if os.environ.get("GROQ_API_KEY"):
        return os.environ["GROQ_API_KEY"]
    # 3. User input in sidebar (fallback)
    return st.session_state.get("api_key", "")


# ─────────────────────────────────────────────
#  CORE FUNCTIONS
# ─────────────────────────────────────────────
def generate_pins(api_key: str, niche: str, num_pins: int, keywords: str) -> list[dict]:
    client = Groq(api_key=api_key)

    kw_note = f"\nFocus keywords to include: {keywords}" if keywords else ""

    prompt = f"""You are an expert Pinterest content strategist.
Generate exactly {num_pins} Pinterest pin ideas for the niche: "{niche}".{kw_note}

Rules:
- Each pin must have a TITLE (max 100 chars, punchy, keyword-rich)
- Each pin must have a DESCRIPTION (150-200 chars, includes a call to action, 3-5 hashtags)
- Vary the content types: tips, how-tos, lists, quotes, product ideas, before/after, inspirational
- Make titles SEO-friendly for Pinterest search
- Descriptions must sound natural, not robotic

Return ONLY a valid JSON array, no extra text. Format:
[
  {{
    "pin_number": 1,
    "title": "...",
    "description": "...",
    "content_type": "tip|how-to|list|quote|product|inspiration"
  }}
]"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.8,
        max_tokens=8000,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def build_docx_bytes(pins: list[dict], niche: str, client_name: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    PINTEREST_RED = RGBColor(0xE6, 0x00, 0x23)
    DARK_GRAY     = RGBColor(0x33, 0x33, 0x33)
    MID_GRAY      = RGBColor(0x77, 0x77, 0x77)

    # Header
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Pinterest Pin Pack")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = PINTEREST_RED

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(f"Niche: {niche.title()}")
    sr.font.size = Pt(13); sr.font.color.rgb = MID_GRAY

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(
        f"Prepared for: {client_name}   |   Pins: {len(pins)}   |   {datetime.now().strftime('%B %d, %Y')}"
    )
    mr.font.size = Pt(10); mr.font.color.rgb = MID_GRAY

    doc.add_paragraph()
    div = doc.add_paragraph("─" * 60)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.runs[0].font.color.rgb = PINTEREST_RED
    doc.add_paragraph()

    # Instructions
    u = doc.add_paragraph()
    ur = u.add_run("HOW TO USE THIS PACK")
    ur.bold = True; ur.font.size = Pt(11); ur.font.color.rgb = PINTEREST_RED

    for inst in [
        "Copy each title into your Pinterest pin title field.",
        "Copy the description into the pin description field.",
        "Add your image, then publish or schedule.",
        "Post 3–5 pins/day for best reach. Mix content types.",
        "Add your website/blog link to every pin!",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(inst).font.size = Pt(10)

    doc.add_paragraph()

    type_colors = {
        "tip": RGBColor(0x1D,0x9E,0x75), "how-to": RGBColor(0x37,0x8A,0xDD),
        "list": RGBColor(0xD8,0x5A,0x30), "quote": RGBColor(0x7F,0x77,0xDD),
        "product": RGBColor(0xBA,0x75,0x17), "inspiration": RGBColor(0xD4,0x53,0x7E),
    }

    for pin in pins:
        num   = pin.get("pin_number", "")
        ctype = pin.get("content_type", "tip").lower()
        color = type_colors.get(ctype, PINTEREST_RED)

        h = doc.add_paragraph()
        nr = h.add_run(f"Pin {num:02d}  "); nr.bold=True; nr.font.size=Pt(11); nr.font.color.rgb=DARK_GRAY
        br = h.add_run(f"[{ctype.upper()}]"); br.bold=True; br.font.size=Pt(9); br.font.color.rgb=color

        tp = doc.add_paragraph()
        tr = tp.add_run(pin.get("title",""))
        tr.bold=True; tr.font.size=Pt(12); tr.font.color.rgb=DARK_GRAY
        tp.paragraph_format.space_before=Pt(2)

        dp = doc.add_paragraph()
        dr = dp.add_run(pin.get("description",""))
        dr.font.size=Pt(10); dr.font.color.rgb=MID_GRAY
        dp.paragraph_format.space_after=Pt(8)

        if num % 5 == 0 and num < len(pins):
            sep = doc.add_paragraph("· " * 30)
            sep.runs[0].font.size = Pt(8)
            sep.runs[0].font.color.rgb = RGBColor(0xDD,0xDD,0xDD)

    doc.add_paragraph()
    fd = doc.add_paragraph("─" * 60)
    fd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fd.runs[0].font.color.rgb = MID_GRAY

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Generated with love. Powered by AI.\nFor more packs, find us on Fiverr.")
    fr.font.size = Pt(9); fr.font.color.rgb = MID_GRAY; fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
#  UI
# ─────────────────────────────────────────────
st.title("📌 Pinterest Pin Pack Generator")
st.markdown("*Generate 50 SEO-optimized Pinterest pins for any niche in seconds.*")
st.divider()

# Sidebar — API key input if not set
with st.sidebar:
    st.header("Settings")
    if not get_api_key():
        api_input = st.text_input(
            "Groq API Key",
            type="password",
            placeholder="gsk_...",
            help="Free at console.groq.com",
        )
        if api_input:
            st.session_state["api_key"] = api_input
            st.success("API key saved!")
    else:
        st.success("API key loaded!")

    st.markdown("---")
    st.markdown("**How it works**")
    st.markdown("1. Enter your niche below\n2. Click Generate\n3. Download your Word doc\n4. Deliver to client!")
    st.markdown("---")
    st.markdown("**Pricing tips**")
    st.markdown("- 25 pins → $10\n- 50 pins → $18\n- 100 pins → $35")

# Main form
col1, col2 = st.columns(2)

with col1:
    niche = st.text_input(
        "Niche *",
        placeholder="e.g. home organization, fitness, budget travel",
    )

with col2:
    client_name = st.text_input(
        "Client name",
        placeholder="e.g. Sarah (for the doc header)",
        value="Client",
    )

keywords = st.text_input(
    "Focus keywords (optional)",
    placeholder="e.g. small apartment, minimalist, DIY",
)

num_pins = st.select_slider(
    "Number of pins",
    options=[25, 50, 75, 100],
    value=50,
)

st.markdown("")
generate_clicked = st.button("Generate Pin Pack", use_container_width=True)

if generate_clicked:
    api_key = get_api_key()
    if not api_key:
        st.error("Please enter your Groq API key in the sidebar first.")
    elif not niche.strip():
        st.error("Please enter a niche.")
    else:
        with st.spinner(f"Generating {num_pins} pins for '{niche}'... (15–30 seconds)"):
            try:
                pins = generate_pins(api_key, niche.strip(), num_pins, keywords.strip())

                c1, c2, c3 = st.columns(3)
                c1.metric("Pins generated", len(pins))
                c2.metric("Content types", len(set(p.get("content_type","") for p in pins)))
                c3.metric("Estimated words", sum(len(p.get("description","").split()) for p in pins))

                docx_bytes = build_docx_bytes(pins, niche.strip(), client_name.strip())
                safe_niche = niche.lower().replace(" ", "_")
                filename   = f"pinterest_pack_{safe_niche}_{datetime.now().strftime('%Y%m%d')}.docx"

                st.markdown('<div class="download-box">', unsafe_allow_html=True)
                st.success("Pack ready! Click below to download.")
                st.download_button(
                    label="⬇️  Download Word Document",
                    data=docx_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                st.markdown('</div>', unsafe_allow_html=True)

                # Preview first 5 pins
                st.markdown("### Preview — first 5 pins")
                for pin in pins[:5]:
                    with st.expander(f"Pin {pin['pin_number']:02d} · {pin.get('content_type','').upper()} · {pin['title'][:60]}..."):
                        st.markdown(f"**Title:** {pin['title']}")
                        st.markdown(f"**Description:** {pin['description']}")

            except json.JSONDecodeError:
                st.error("AI returned unexpected format. Please try again.")
            except Exception as e:
                st.error(f"Error: {e}")

st.markdown("---")
st.caption("Powered by Groq (Llama 3.3) · Built for Fiverr sellers · Zero subscription needed")
